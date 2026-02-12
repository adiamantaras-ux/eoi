from __future__ import annotations

from pathlib import Path
from datetime import date
import re

from django.conf import settings
from django.core.management.base import BaseCommand, CommandError
from django.utils import timezone

try:
    from docx import Document  # python-docx
except ModuleNotFoundError as e:
    raise ModuleNotFoundError(
        "Λείπει η βιβλιοθήκη python-docx. Τρέξε: pip install python-docx"
    ) from e

from registry.models import Athlete


def _fmt_date(d: date | None) -> str:
    if not d:
        return ""
    # 29/12/2025
    return d.strftime("%d/%m/%Y")


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """
    Αντικαθιστά μέσα στο paragraph.
    Χρησιμοποιούμε paragraph.text για απλότητα (χάνει formatting, αλλά το πρότυπο είναι απλό).
    """
    txt = paragraph.text or ""
    original = txt

    for pattern, value in replacements.items():
        txt = re.sub(pattern, value, txt, flags=re.IGNORECASE)

    if txt != original:
        paragraph.text = txt


def _replace_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    # paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)


class Command(BaseCommand):
    help = "Δημιουργεί Word ιατρικής βεβαίωσης (συμπληρωμένο) για έναν αθλητή."

    def add_arguments(self, parser):
        parser.add_argument("athlete_id", type=int, help="ID αθλητή (pk)")
        parser.add_argument(
            "--out",
            type=str,
            default="",
            help="Προαιρετικά: custom φάκελος εξόδου (αλλιώς πάει στο MEDIA_ROOT).",
        )

    def handle(self, *args, **options):
        athlete_id: int = options["athlete_id"]
        out_override: str = options["out"] or ""

        athlete = Athlete.objects.filter(pk=athlete_id).select_related("club").first()
        if not athlete:
            raise CommandError(f"Δεν βρέθηκε αθλητής με id={athlete_id}")

        # ---- template path (from settings) ----
        template_path = getattr(settings, "MEDICAL_CERT_TEMPLATE_PATH", None)
        if not template_path:
            raise CommandError("Λείπει το MEDICAL_CERT_TEMPLATE_PATH από το settings.py")

        template_path = Path(template_path)
        if not template_path.exists():
            raise CommandError(f"Δεν βρέθηκε το πρότυπο αρχείο: {template_path}")

        # ---- output root ----
        if out_override:
            out_root = Path(out_override)
        else:
            media_root = getattr(settings, "MEDIA_ROOT", Path(settings.BASE_DIR) / "media")
            out_root = Path(media_root)

        # ---- output folder per athlete ----
        out_dir = out_root / "medical_certificates" / str(athlete.pk)
        out_dir.mkdir(parents=True, exist_ok=True)

        today = timezone.localdate()

        # ασφαλές όνομα αρχείου
        reg = (athlete.eoi_registry_number or f"ID{athlete.pk}").strip()
        reg_safe = re.sub(r"[^\w\-]+", "_", reg, flags=re.UNICODE)

        filename = f"MedicalCertificate_{reg_safe}_{today.isoformat()}.docx"
        out_path = out_dir / filename

        # ---- load doc ----
        doc = Document(str(template_path))

        full_name = f"{(athlete.last_name or '').strip()} {(athlete.first_name or '').strip()}".strip()
        club_name = ""
        if athlete.club:
            # π.χ. "ΙΟΠ - Ιππικός Όμιλος ..."
            club_name = str(athlete.club)

        # Εδώ κάνουμε replace “με βάση τις ετικέτες” που ήδη έχει το πρότυπο (χωρίς να το αλλάξεις)
        # Πιάνει γραμμές τύπου: "ΟΝΟΜΑΤΕΠΩΝΥΜΟ: ...."
        replacements = {
            r"(ΟΝΟΜΑΤΕΠΩΝΥΜΟ\s*:\s*)(.*)$": r"\1" + (full_name or ""),
            r"(ΗΜΕΡΟΜΗΝΙΑ\s+ΓΕΝΝΗΣΗΣ\s*)(\.*\s*)": r"\1" + _fmt_date(getattr(athlete, "birth_date", None)),
            r"(ΑΡΙΘΜΟΣ\s+ΤΑΥΤΟΤΗΤΑΣ\s*.*?:\s*)(.*)$": r"\1" + (athlete.id_number or ""),
            r"(ΤΗΛ\.\s*:\s*)(.*)$": r"\1" + "",
            r"(Email\s*:\s*)(.*)$": r"\1" + (getattr(athlete, "email", "") or ""),
            r"(ΣΩΜΑΤΕΙΟ\s*)(\.*\s*)": r"\1" + (club_name or ""),
        }

        _replace_everywhere(doc, replacements)

        doc.save(str(out_path))

        self.stdout.write(self.style.SUCCESS("OK"))
        self.stdout.write(f"Αθλητής: {athlete}")
        self.stdout.write(f"Template: {template_path}")
        self.stdout.write(f"Έξοδος: {out_path}")
